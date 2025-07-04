from fastapi import FastAPI, HTTPException, Form, UploadFile, File, Query, Request
from fastapi.responses import JSONResponse, RedirectResponse
from pydantic import BaseModel
from sentence_transformers import SentenceTransformer, util
import google.generativeai as genai
from nameparser import HumanName
from typing import List
from docx import Document
import io
import urllib.parse
import json
from dotenv import load_dotenv
from fastapi.middleware.cors import CORSMiddleware
from urllib.parse import quote
import httpx
import uvicorn
import fitz
import docx
import openai
import os
import hashlib
import csv
import re
from datetime import datetime
from dateutil import parser
import spacy
import pdfplumber
import docx
import logging
import requests
from PIL import Image
from typing import Dict, List, Optional
from models import Base, engine

app = FastAPI()

from chatbot_module import router as chatbot_router
app.include_router(chatbot_router)
Base.metadata.create_all(bind=engine)

model = SentenceTransformer('all-MiniLM-L6-v2')
nlp = spacy.load("en_core_web_trf")

PROCESSED_CVS_FILE = "processed_cvs.json"

if os.path.exists(PROCESSED_CVS_FILE):
    with open(PROCESSED_CVS_FILE, "r") as f:
        processed_resumes = json.load(f)
else:
    processed_resumes = {}

load_dotenv()
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
genai.configure(api_key=GEMINI_API_KEY)

ROLE_SYNONYMS = {
    "software developer": "software engineer",
    "web developer": "frontend engineer",
    "data scientist": "machine learning engineer"
}

MONTHS_MAPPING = {
    'jan': '01', 'january': '01',
    'feb': '02', 'february': '02',
    'mar': '03', 'march': '03',
    'apr': '04', 'april': '04',
    'may': '05',
    'jun': '06', 'june': '06',
    'jul': '07', 'july': '07',
    'aug': '08', 'august': '08',
    'sep': '09', 'september': '09',
    'oct': '10', 'october': '10',
    'nov': '11', 'november': '11',
    'dec': '12', 'december': '12',
}

openai.api_key = os.getenv("OPENAI_API_KEY")

def extract_details_using_gpt(text: str) -> dict:
    prompt = f"""
Extract the following details from this resume strictly in JSON format:
- Full Name
- Email
- Mobile Number
- Total Years of Experience (based on years in job history)
- Skills (as a list)

Resume:
\"\"\"
{text}
\"\"\"
Return JSON only with keys: full_name, email, mobile_number, total_experience, skills.
If any field is missing, return null or an empty list.
"""

    try:
        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            temperature=0,
            messages=[{"role": "user", "content": prompt}],
        )
        content = response["choices"][0]["message"]["content"]
        # Try parsing GPT JSON safely
        import json
        match = re.search(r'\{.*\}', content, re.DOTALL)
        if match:
            return json.loads(match.group())
        else:
            return {"error": "Failed to parse JSON from GPT", "raw_output": content}
    except openai.error.OpenAIError as e:
        return {"error": str(e)}
    
def clean_text(text: str) -> str:
    return re.sub(r'\n+', '\n', text.strip())

def call_gpt(text: str) -> dict:
    system_prompt = "You are a resume parser. Extract full name, email, mobile number, total years of experience, and skills as a JSON."
    user_prompt = f"Extract information from the following resume text:\n\n{text[:4000]}"  # Keep input within token limit

    response = openai.chat.completions.create(
        model="gpt-3.5-turbo",
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt}
        ],
        temperature=0.2
    )

    content = response.choices[0].message.content
    try:
        # Expecting valid JSON in the response
        import json
        return json.loads(content)
    except:
        return {"raw_output": content, "error": "Failed to parse JSON from GPT"}
    
    
def normalize_role(role: str) -> str:
    role_lower = role.lower().strip()
    return ROLE_SYNONYMS.get(role_lower, role_lower)

def extract_text_from_uploaded_file(file: UploadFile) -> str:
    filename = file.filename.lower()
    if filename.endswith(".pdf"):
        with pdfplumber.open(file.file) as pdf:
            return "\n".join([page.extract_text() or "" for page in pdf.pages])
    elif filename.endswith(".docx"):
        document = docx.Document(file.file)
        return "\n".join([para.text for para in document.paragraphs])
    else:
        raise ValueError("Unsupported file type")

def extract_years(text: str) -> float:
    text = text.lower()

    # Common phrasing variations
    experience_patterns = [
        r'(\d{1,2}(?:\.\d+)?)\s*\+?\s*(?:years|yrs)\s+(?:of\s+)?(?:experience|exp)',
        r'(?:experience|exp)\s+(?:of\s+)?(\d{1,2}(?:\.\d+)?)\s*\+?\s*(?:years|yrs)',
        r'(\d{1,2}(?:\.\d+)?)\s*(?:years|yrs)',
        r'(\d{1,2})\s*\+?\s*(?:years|yrs)\s*experience'
    ]

    max_years = 0.0

    for pattern in experience_patterns:
        matches = re.findall(pattern, text)
        for match in matches:
            try:
                # Handle match coming as tuple (grouped regex)
                years = float(match[0] if isinstance(match, tuple) else match)
                if 0 < years < 40:  # filter out unrealistic numbers
                    max_years = max(max_years, years)
            except:
                continue

    return round(max_years, 2)

def extract_experience_from_dates(text: str) -> float:
    text = text.lower()
    now = datetime.now()

    # Patterns to match various job date formats
    patterns = [
        r'([a-z]{3,9})[\s\-]*(\d{4})\s*(?:to|–|[-])\s*(present|current|[a-z]{3,9}[\s\-]*\d{4})',
        r'(\d{4})\s*(?:to|–|[-])\s*(present|current|\d{4})'
    ]

    total_months = 0

    for pattern in patterns:
        for match in re.findall(pattern, text):
            try:
                if len(match) == 3:  # e.g., Jan 2020 - Present
                    start_month, start_year, end = match
                    start_month = MONTHS_MAPPING.get(start_month[:3], '01')
                    start_date = parser.parse(f"{start_year}-{start_month}-01")

                elif len(match) == 2:  # e.g., 2018 - 2021
                    start_year, end = match
                    start_date = parser.parse(f"{start_year}-01-01")

                if 'present' in end or 'current' in end:
                    end_date = now
                else:
                    try:
                        if re.match(r'[a-z]{3,9}', end):
                            end_month, end_year = end.split()
                            end_month = MONTHS_MAPPING.get(end_month[:3], '01')
                        elif re.match(r'[a-z]{3,9}\s*\d{4}', end):
                            end_month, end_year = end[:3], end[-4:]
                            end_month = MONTHS_MAPPING.get(end_month[:3], '01')
                        else:
                            end_date = parser.parse(f"{end}-01-01")
                            raise Exception()  # Skip to except
                        end_date = parser.parse(f"{end_year}-{end_month}-01")
                    except:
                        end_date = parser.parse(f"{end}-01-01")

                months = (end_date.year - start_date.year) * 12 + (end_date.month - start_date.month)
                if 0 < months < 600:
                    total_months += months
            except Exception as e:
                continue

    return round(total_months / 12, 2) if total_months > 0 else 0.0

def extract_text_from_pdf(file_path: str) -> str:
    try:
        with pdfplumber.open(file_path) as pdf:
            text = "\n".join([page.extract_text() or "" for page in pdf.pages])
            return text.strip()
    except:
        return ""

def extract_text_from_docx(file_path: str) -> str:
    try:
        doc = docx.Document(file_path)
        return "\n".join(para.text for para in doc.paragraphs).strip()
    except:
        return ""
    
def get_combined_hash(file_path: str, job_description: str) -> str:
    with open(file_path, "rb") as f:
        file_content = f.read()
    combined = file_content + job_description.encode('utf-8')
    return hashlib.md5(combined).hexdigest()

def clean_text_for_hashing(text: str) -> str:
    # Remove emails
    text = re.sub(r'\b[\w.-]+?@\w+?\.\w+?\b', '', text)
    # Remove phone numbers
    text = re.sub(r'(\+?\d{1,3}[-.\s]?)?(\d{10}|\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4})', '', text)
    # Lowercase and remove extra spaces
    text = text.lower().strip().replace('\n', ' ')
    return text

logging.getLogger("pdfminer").setLevel(logging.ERROR)

def get_text_hash(file_path: str) -> str:
    if file_path.lower().endswith(".pdf"):
        text = extract_text_from_pdf(file_path)
    elif file_path.lower().endswith(".docx"):
        text = extract_text_from_docx(file_path)
    else:
        return None

    cleaned_text = clean_text_for_hashing(text)
    return hashlib.md5(cleaned_text.encode('utf-8')).hexdigest()

def parse_experience_with_chatgpt(text: str) -> float:
    prompt = f"""
You are a resume parser. Extract only the total years of professional work experience from the following resume. Do not count internships or academic projects. If the resume mentions durations like "March 2020 - Present" or "3+ years experience in development", calculate the total accurately.

Resume:
{text}

Return only the total years as a float (e.g., 2.5 or 4.0). If unclear, return 0.
"""
    try:
        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            temperature=0,
            messages=[{"role": "user", "content": prompt}]
        )
        result = response.choices[0].message.content.strip()
        match = re.search(r'(\d+(\.\d+)?)', result)
        if match:
            return round(float(match.group(1)), 2)
    except Exception:
        pass
    return 0.0

def extract_text_from_path(path: str) -> str:
    if path.lower().endswith(".pdf"):
        return extract_text_from_pdf(path)
    elif path.lower().endswith(".docx"):
        return extract_text_from_docx(path)
    return ""

def extract_text(file: UploadFile) -> str:
    if file.filename.endswith(".pdf"):
        with pdfplumber.open(file.file) as pdf:
            return "\n".join(page.extract_text() or "" for page in pdf.pages)
    elif file.filename.endswith(".docx"):
        doc = docx.Document(file.file)
        return "\n".join([para.text for para in doc.paragraphs])
    else:
        raise ValueError("Only PDF and DOCX files are supported")

def extract_email(text):
    match = re.search(r'\b[\w.-]+?@\w+?\.\w+?\b', text)
    return match.group(0) if match else None

def extract_phone(text):
    match = re.search(r'(\+?\d{1,3}[-.\s]?)?(\d{10})', text)
    return match.group(0) if match else None

def extract_name(text):
    lines = text.strip().split('\n')
    for line in lines[:10]:  # First 10 lines are likely to contain the name
        if len(line.strip().split()) in [2, 3] and line[0].isupper():
            doc = nlp(line.strip())
            for ent in doc.ents:
                if ent.label_ == "PERSON":
                    return ent.text
    return "Unknown Name"

def extract_skills(text):
    skill_keywords = [
        'python', 'java', 'c++', 'javascript', 'react', 'node', 'angular',
        'sql', 'mongodb', 'docker', 'kubernetes', 'aws', 'gcp', 'azure',
        'html', 'css', 'tensorflow', 'pytorch', 'flask', 'fastapi'
    ]
    text_lower = text.lower()
    found = [skill for skill in skill_keywords if skill in text_lower]
    return list(set(found))

def normalize_years(text: str) -> str:
    return re.sub(r'(?:(?:\d\s*){4})', lambda m: ''.join(m.group(0).split()), text)

# Utility: Extract text from PDF/DOCX
def extract_text_from_path(filepath: str) -> str:
    try:
        if filepath.endswith(".pdf"):
            # Try pdfplumber first
            with pdfplumber.open(filepath) as pdf:
                text = "\n".join(page.extract_text() or "" for page in pdf.pages)
            if len(text.strip()) > 300:
                return text
            # fallback to fitz
            doc = fitz.open(filepath)
            return "\n".join([page.get_text() for page in doc])
        elif filepath.endswith(".docx"):
            doc = docx.Document(filepath)
            return "\n".join([p.text for p in doc.paragraphs])
        else:
            return ""
    except:
        return ""

# Utility: Parse month-year to datetime
def parse_date(month_str: str, year_str: str) -> Optional[datetime]:
    try:
        month = MONTHS_MAPPING.get(month_str[:3].lower(), '01')
        return parser.parse(f"{year_str}-{month}-01")
    except:
        return None

def extract_experience(text: str) -> float:
    text = normalize_years(text.lower())
    now = datetime.now()
    total_months = 0
    seen_ranges = set()
    job_titles = ["engineer", "developer", "consultant", "intern", "analyst", "manager", "associate", "lead", "designer", "architect", "specialist", "executive"]
    lines = text.split('\n')

    for idx, line in enumerate(lines):
        if any(title in line.lower() for title in job_titles):
            context = " ".join(lines[max(0, idx - 2):idx + 5])
            date_patterns = re.findall(r'([a-z]{3,9})[\s\-]*(\d{4})\s*(?:to|–|[-])\s*([a-z]{3,9}|present|current)[\s\-]*(\d{4})?', context, re.IGNORECASE)
            for sm, sy, em, ey in date_patterns:
                try:
                    start_date = parse_date(sm, sy)
                    end_date = now if em in ['present', 'current'] else parse_date(em, ey or sy)
                    if not start_date or not end_date or end_date < start_date:
                        continue
                    if (start_date, end_date) in seen_ranges:
                        continue
                    seen_ranges.add((start_date, end_date))
                    months = (end_date.year - start_date.year) * 12 + (end_date.month - start_date.month)
                    if 0 < months <= 480:
                        total_months += months
                except:
                    continue

    textual_matches = re.findall(r'(\d+(?:\.\d+)?)\s*\+?\s*(?:years|yrs|year)\s*(?:of)?\s*(?:experience|exp)?', text)
    max_textual_years = max([float(y) for y in textual_matches if float(y) <= 50], default=0.0)
    inferred_years = round(total_months / 12, 2)
    if 0.5 <= inferred_years <= 50 and abs(inferred_years - max_textual_years) <= 2:
        return inferred_years
    return max(inferred_years, max_textual_years)

# ============ GPT Scoring ============

def gpt_relevance_score(resume_text: str, job_description: str) -> Optional[float]:
    prompt = f"""
You are a resume evaluator. Given the resume and the job description, rate how well the resume matches the job description on a scale from 0 to 100.

Scoring Criteria:
- Relevant experience and skills mentioned
- Education or certifications matching the job
- Alignment in roles, tools, technologies, and industry
- Communication and structure quality of the resume

Only respond with a single number between 0 and 100 — no words, no explanation, no symbols.

Resume:
\"\"\"{resume_text[:4000]}\"\"\"

Job Description:
\"\"\"{job_description[:2000]}\"\"\"
"""

    try:
        client = openai.OpenAI()
        response = client.chat.completions.create(
            model="gpt-4",
            messages=[
                {"role": "system", "content": "You are an expert recruiter evaluating resume-job fit scores."},
                {"role": "user", "content": prompt}
            ]
        )
        score_str = response.choices[0].message.content.strip()
        match = re.search(r'\d{1,3}(?:\.\d+)?', score_str)
        if match:
            score = float(match.group())
            return max(0.0, min(score, 100.0))  # Clamp between 0 and 100
    except Exception as e:
        print(f"Error: {e}")
        return 0.0

    return 0.0

def calculate_experience_bonus(required: float, actual: float) -> float:
    try:
        required, actual = float(required), float(actual)
    except:
        return 0.0
    if actual >= required:
        return 1.0
    elif required == 0:
        return 0.0
    return round(actual / required, 2)

def get_combined_hash(file_path: str, job_description: str) -> str:
    with open(file_path, "rb") as f:
        content = f.read()
    return hashlib.md5(content + job_description.encode()).hexdigest()

@app.post("/generate-jd")
def generate_job_description(job_title: str = Form(...), skills: str = Form(...), experience: str = Form(...)):
    prompt = (
        f"Generate a professional job description for a {job_title} role.\n"
        f"Required experience: {experience} years.\n"
        f"Skills: {skills}.\n"
        f"Format the output in bullet points under responsibilities and requirements.\n"
        f"return job description only in HTML Format and points should be in <ul><li>\n"
        f"do not give job description heading location, and experience as only return JD, Key Responsibilities and Requirements.\n"
    )

    try:
        model = genai.GenerativeModel("gemini-1.5-flash-latest")
        response = model.generate_content(prompt)
        return {"job_description": response.text}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/job-embed")
def get_job_embedding(
    job_title: str = Form(...),
    experience: str = Form(...),
    skills: str = Form(...),
    description: str = Form(...)
):
    norm_job_title = normalize_role(job_title)
    parsed_exp = extract_years(experience)
    full_text = f"{norm_job_title} {parsed_exp} years experience {skills} {description}"
    embedding = model.encode(full_text).tolist()

    row = {
        "timestamp": datetime.now().isoformat(),
        "job_title": job_title,
        "experience": experience,
        "skills": skills,
        "description": description,
        "embedding": json.dumps(embedding, indent=2)
    }

    file_exists = os.path.exists("job_embeddings.csv")
    with open("job_embeddings.csv", mode="a", newline='', encoding='utf-8') as csvfile:
        fieldnames = ["timestamp", "job_title", "experience", "skills", "description", "embedding"]
        writer = csv.DictWriter(csvfile, fieldnames=fieldnames)
        if not file_exists:
            writer.writeheader()
        writer.writerow(row)

    return {"embedding": embedding}

@app.post("/match-resumes")
def match_resumes(
    job_description: str = Form(...),
    resume_folder_path: str = Form(...),
    required_experience: str = Form("0")
):
    if not os.path.exists(resume_folder_path):
        raise HTTPException(status_code=400, detail="Resume folder path does not exist.")

    results = []

    for filename in os.listdir(resume_folder_path):
        filepath = os.path.join(resume_folder_path, filename)
        if not os.path.isfile(filepath) or not filename.lower().endswith((".pdf", ".docx")):
            continue

        # Extract text
        text = extract_text_from_path(filepath)

        # Extract actual experience
        actual_exp = extract_experience(text)

        # Calculate bonus
        exp_bonus = calculate_experience_bonus(required_experience, actual_exp)

        # Get GPT relevance score
        relevance = gpt_relevance_score(text, job_description)

        # Save
        score_data = {
            "filename": filename,
            "actual_experience_years": round(actual_exp, 2),
            "experience_bonus": round(exp_bonus, 2),
            "relevance_score": round(relevance or 0.0, 2)
        }

        results.append(score_data)

    return results

#parse resume
@app.post("/parse-resume")
async def parse_resume(file: UploadFile = File(...)):
    try:
        text = extract_text(file)
        gpt_result = call_gpt(text)
        return gpt_result
    except Exception as e:
        return {"error": str(e)}

#find duplicates
@app.post("/find-duplicate-resumes")
def find_duplicates(resume_folder_path: str = Form(...)):
    if not os.path.exists(resume_folder_path):
        raise HTTPException(status_code=400, detail="Resume folder path does not exist.")

    hash_map: Dict[str, List[str]] = {}

    for filename in os.listdir(resume_folder_path):
        filepath = os.path.join(resume_folder_path, filename)
        if not os.path.isfile(filepath):
            continue

        hash_val = get_text_hash(filepath)
        if hash_val is None:
            continue

        hash_map.setdefault(hash_val, []).append(filename)

    # Filter to only include hashes with duplicates
    duplicates = {h: f for h, f in hash_map.items() if len(f) > 1}

    return {"duplicates": duplicates, "total_duplicates": len(duplicates)}

if __name__ == "__main__":
    uvicorn.run(app, host="127.0.0.1", port=8000, reload=True)
    